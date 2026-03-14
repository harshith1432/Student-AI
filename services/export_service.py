import os
import markdown
from bs4 import BeautifulSoup
from xhtml2pdf import pisa
from io import BytesIO
from docx import Document
from datetime import datetime

# Path where we'll store static files temporarily if needed, or we just transmit bytes.
# We will use BytesIO to directly stream to the user without saving on disk.

def convert_markdown_to_html(md_text):
    """Convert markdown to HTML"""
    return markdown.markdown(md_text, extensions=['extra', 'nl2br'])

def generate_pdf(text, is_handwritten=False):
    """
    Generate a PDF from text. If is_handwritten is True, apply a handwritten font.
    Using xhtml2pdf for simplicity of rendering HTML to PDF.
    """
    html_content = convert_markdown_to_html(text)
    
    font_style = ""
    # We create a simple CSS inline to handle the handwritten font
    # Note: Using standard fonts or a web font for demo. For real handwriting, 
    # we would need to point to a local TTF file or a Google Font (like 'Caveat').
    if is_handwritten:
        # Using a generic cursive font and removing network import to avoid xhtml2pdf issues
        font_style = """
        body {
            font-family: 'Cursive', 'Comic Sans MS', cursive;
            font-size: 16pt;
            background-color: #F0F8FF;
            color: #000080;
            line-height: 1.4;
        }
        """
    else:
        font_style = """
        body {
            font-family: Helvetica, Arial, sans-serif;
            font-size: 12pt;
            line-height: 1.5;
            color: #1a1a1a;
        }
        """

    full_html = f"""
    <html>
    <head>
        <style>
            {font_style}
            h1, h2, h3 {{ color: #1a1a1a; }}
            p {{ margin-bottom: 15px; }}
            ul, ol {{ margin-bottom: 15px; }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """

    pdf_file = BytesIO()
    pisa.CreatePDF(BytesIO(full_html.encode('utf-8')), dest=pdf_file)
    pdf_file.seek(0)
    return pdf_file

def generate_docx(text):
    """
    Generate a Word document from text.
    For simplicity, we strip HTML from the converted markdown and add paragraphs.
    """
    html_content = convert_markdown_to_html(text)
    soup = BeautifulSoup(html_content, 'html.parser')
    
    document = Document()
    document.add_heading('Student AI Document', 0)
    
    # Very basic parsing of HTML to Word
    # In a real app we'd map h1->Heading1, etc.
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'li', 'ol']):
        if element.name in ['h1', 'h2', 'h3']:
            level = int(element.name[1])
            document.add_heading(element.get_text(), level=level)
        elif element.name == 'p':
            document.add_paragraph(element.get_text())
        elif element.name == 'li':
            document.add_paragraph(element.get_text(), style='List Bullet')
            
    docx_file = BytesIO()
    document.save(docx_file)
    docx_file.seek(0)
    return docx_file

def generate_txt(text):
    """
    Generate a simple fast text file format from markdown
    """
    txt_file = BytesIO()
    txt_file.write(text.encode('utf-8'))
    txt_file.seek(0)
    return txt_file
